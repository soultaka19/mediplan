# -*- coding: utf-8 -*-
"""Convertit le cahier des charges Markdown en .docx lisible.

Volontairement minimal : titres, tableaux, listes, citations, gras et code.
Suffisant pour un document de remise ; pas un moteur Markdown complet.
"""
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

SRC, DST = sys.argv[1], sys.argv[2]
FOOTER = sys.argv[3] if len(sys.argv) > 3 else "MediPlan — Collège La Cité — Printemps 2026"

ACCENT = RGBColor(0x1E, 0x5F, 0xA8)
GREY = RGBColor(0x55, 0x5F, 0x6D)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)

for lvl, size in ((1, 20), (2, 15), (3, 12.5), (4, 11)):
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = "Calibri"
    h.font.size = Pt(size)
    h.font.color.rgb = ACCENT
    h.font.bold = True


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
IMAGE = re.compile(r"!\[([^\]]*)\]\(([^)\s]+)\)")
SRC_DIR = os.path.dirname(os.path.abspath(SRC))


def add_image(par, path, alt, width_in):
    """Insère une image ; si le fichier manque, laisse le texte alternatif."""
    full = os.path.normpath(os.path.join(SRC_DIR, path))
    if os.path.exists(full):
        par.add_run().add_picture(full, width=Inches(width_in))
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        par.add_run(f"[image absente : {alt or path}]").italic = True


def add_runs(par, text, img_width=6.2):
    """Découpe le gras, le code inline et les images ; le reste passe tel quel."""
    # Les images sont traitées d'abord : elles ne se mélangent pas au texte.
    if IMAGE.search(text):
        pos = 0
        for m in IMAGE.finditer(text):
            before = text[pos : m.start()].strip()
            if before:
                add_runs(par, before, img_width)
            add_image(par, m.group(2), m.group(1), img_width)
            pos = m.end()
        rest = text[pos:].strip()
        if rest:
            add_runs(par, rest, img_width)
        return
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            par.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = par.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        else:
            par.add_run(chunk)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


# Les emoji font basculer Word sur Segoe UI Emoji, qu'il embarque en entier :
# 4 Mo de PDF pour un document de texte. Ici ils sont decoratifs — le mot qui
# suit porte le sens — donc on les retire a la conversion. Le Markdown source
# les garde : ils aident a la lecture sur GitHub.
EMOJI = re.compile(
    "[←-⇿⌀-⏿①-➿⬀-⯿️‍]"
    "|[\U0001f000-\U0001faff]"
)


def strip_emoji(text):
    return re.sub(r" {2,}", " ", EMOJI.sub("", text)).strip()


lines = [
    strip_emoji(l) if EMOJI.search(l) else l
    for l in open(SRC, encoding="utf-8").read().split("\n")
]
i = 0
while i < len(lines):
    ln = lines[i]
    stripped = ln.strip()

    # --- Tableau : une ligne | … | suivie d'un séparateur |---|
    if (
        stripped.startswith("|")
        and i + 1 < len(lines)
        and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip())
    ):
        header = split_row(stripped)
        i += 2
        body = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            body.append(split_row(lines[i].strip()))
            i += 1
        t = doc.add_table(rows=1, cols=len(header))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c, txt in zip(t.rows[0].cells, header):
            c.paragraphs[0].clear()
            add_runs(c.paragraphs[0], txt, 3.0)
            for r in c.paragraphs[0].runs:
                r.bold = True
        for row in body:
            cells = t.add_row().cells
            for c, txt in zip(cells, row[: len(header)]):
                c.paragraphs[0].clear()
                add_runs(c.paragraphs[0], txt, 3.0)
        doc.add_paragraph()
        continue

    # --- Bloc de code
    if stripped.startswith("```"):
        i += 1
        buf = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            buf.append(lines[i])
            i += 1
        i += 1
        p = doc.add_paragraph()
        r = p.add_run("\n".join(buf))
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        continue

    if not stripped:
        i += 1
        continue

    # --- Filet horizontal : saut de section
    if re.match(r"^-{3,}$", stripped):
        doc.add_paragraph()
        i += 1
        continue

    # --- Titres
    m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
    if m:
        lvl, txt = len(m.group(1)), m.group(2)
        p = doc.add_heading(level=lvl)
        add_runs(p, txt)
        i += 1
        continue

    # --- Citation
    if stripped.startswith(">"):
        txt = re.sub(r"^>\s?", "", stripped)
        i += 1
        while i < len(lines) and lines[i].strip().startswith(">"):
            nxt = re.sub(r"^>\s?", "", lines[i].strip())
            txt = (txt + "\n") if not nxt else (txt + " " + nxt)
            i += 1
        p = doc.add_paragraph(style="Intense Quote")
        add_runs(p, txt.strip())
        continue

    # --- Listes
    m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_runs(p, m.group(2))
        i += 1
        continue
    if re.match(r"^[-*]\s+", stripped):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, re.sub(r"^[-*]\s+", "", stripped))
        i += 1
        continue

    # --- Paragraphe : on agrège les lignes jusqu'à la ligne vide
    buf = [stripped]
    i += 1
    while i < len(lines):
        n = lines[i].strip()
        if not n or n.startswith(("#", ">", "|", "```", "- ", "* ")) or re.match(r"^\d+\.\s", n) or re.match(r"^-{3,}$", n):
            break
        buf.append(n)
        i += 1
    p = doc.add_paragraph()
    add_runs(p, " ".join(buf))

# Pied de page : identification du document
sec = doc.sections[0]
foot = sec.footer.paragraphs[0]
foot.text = FOOTER
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.runs[0].font.size = Pt(8)
foot.runs[0].font.color.rgb = GREY

doc.save(DST)
print("OK ->", DST)
